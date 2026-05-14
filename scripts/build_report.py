from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts"
ASSET_DIR = OUT_DIR / "report_assets"
REPORT_PATH = OUT_DIR / "Отчет_Криптоброкер_Kafka_мессенджер_сенсоры.docx"


def font(size):
    for candidate in [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
    ]:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    f = font(21)
    total_height = len(lines) * 25
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=f, fill=(20, 30, 45))
        y += 25


def draw_arrow(draw, start, end, label):
    color = (68, 84, 106)
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, _ = start
    if ex >= sx:
        arrow = [(ex, ey), (ex - 15, ey - 8), (ex - 15, ey + 8)]
    else:
        arrow = [(ex, ey), (ex + 15, ey - 8), (ex + 15, ey + 8)]
    draw.polygon(arrow, fill=color)
    f = font(17)
    mx = (start[0] + end[0]) / 2
    my = (start[1] + end[1]) / 2 - 28
    bbox = draw.textbbox((0, 0), label, font=f)
    draw.rounded_rectangle(
        (mx - (bbox[2] - bbox[0]) / 2 - 8, my - 4, mx + (bbox[2] - bbox[0]) / 2 + 8, my + 22),
        radius=7,
        fill=(255, 255, 255),
    )
    draw.text((mx - (bbox[2] - bbox[0]) / 2, my), label, font=f, fill=color)


def architecture_png(path):
    img = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "Архитектура: мессенджер и сенсоры через криптоконтейнеры", font=font(34), fill=(11, 37, 69))
    draw.rounded_rectangle((35, 95, 1565, 705), radius=26, outline=(208, 215, 222), width=3, fill=(248, 250, 252))

    boxes = {
        "app": (60, 190, 220, 270, "Web App\nchat", (218, 232, 252), (108, 142, 191)),
        "sensor_gen": (60, 420, 220, 500, "Sensor\nGenerator", (224, 242, 241), (80, 156, 144)),
        "chat_raw": (300, 190, 480, 270, "Kafka\nmessages.raw", (255, 242, 204), (214, 182, 86)),
        "sensor_raw": (300, 420, 480, 500, "Kafka\nsensors.raw", (255, 242, 204), (214, 182, 86)),
        "gateway": (570, 300, 780, 410, "Crypto Gateway\nPEP + Facade", (213, 232, 212), (130, 179, 102)),
        "policy": (570, 120, 780, 210, "Policy Engine\nPDP", (232, 222, 248), (150, 115, 166)),
        "chat_crypto": (870, 190, 1060, 270, "Kafka\nmessages.crypto", (255, 242, 204), (214, 182, 86)),
        "sensor_crypto": (870, 420, 1060, 500, "Kafka\nsensors.crypto", (255, 242, 204), (214, 182, 86)),
        "filter": (1160, 300, 1340, 410, "Filter\ndecrypt + sanitize", (248, 206, 204), (184, 84, 80)),
        "result": (1410, 300, 1540, 410, "Filtered\nUI/Grafana", (255, 230, 204), (215, 155, 0)),
    }
    for box in boxes.values():
        draw_box(draw, box[:4], box[4], box[5], box[6])

    draw_arrow(draw, (220, 230), (300, 230), "chat")
    draw_arrow(draw, (220, 460), (300, 460), "auto sensors")
    draw_arrow(draw, (480, 230), (570, 335), "raw")
    draw_arrow(draw, (480, 460), (570, 370), "raw")
    draw_arrow(draw, (675, 300), (675, 210), "allow/deny")
    draw_arrow(draw, (780, 335), (870, 230), "ciphertext")
    draw_arrow(draw, (780, 370), (870, 460), "ciphertext")
    draw_arrow(draw, (1060, 230), (1160, 335), "decrypt")
    draw_arrow(draw, (1060, 460), (1160, 370), "decrypt")
    draw_arrow(draw, (1340, 355), (1410, 355), "filtered")
    img.save(path)


def sequence_png(path):
    img = Image.new("RGB", (1600, 720), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "Диаграмма взаимодействия", font=font(34), fill=(11, 37, 69))
    xs = [110, 310, 520, 740, 970, 1190, 1430]
    names = ["User", "Web App", "Sensor Gen", "Crypto Gateway", "Policy Engine", "Filter", "UI/Grafana"]
    fills = [(218, 232, 252), (218, 232, 252), (224, 242, 241), (213, 232, 212), (232, 222, 248), (248, 206, 204), (255, 230, 204)]
    outlines = [(108, 142, 191), (108, 142, 191), (80, 156, 144), (130, 179, 102), (150, 115, 166), (184, 84, 80), (215, 155, 0)]
    for x, name, fill, outline in zip(xs, names, fills, outlines):
        draw_box(draw, (x - 95, 115, x + 95, 180), name, fill, outline)
        draw.line((x, 180, x, 610), fill=(208, 215, 222), width=3)
    steps = [
        (xs[0], xs[1], 240, "1. input"),
        (xs[1], xs[3], 300, "2a. message raw"),
        (xs[2], xs[3], 360, "2b. sensor raw"),
        (xs[3], xs[4], 420, "3. policy request"),
        (xs[4], xs[3], 480, "4. allow/deny"),
        (xs[3], xs[5], 540, "5. crypto container"),
        (xs[5], xs[6], 600, "6. filtered topic"),
    ]
    for sx, ex, y, label in steps:
        draw_arrow(draw, (sx, y), (ex, y), label)
    img.save(path)


def classes_png(path):
    img = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "Диаграмма классов и шаблоны", font=font(34), fill=(11, 37, 69))
    boxes = [
        (70, 145, 420, 305, "MessageTypeStrategy\n+ detect(payload, topic)\nStrategy", (218, 232, 252), (108, 142, 191)),
        (485, 145, 835, 305, "PolicyDecisionPoint\n+ decide(request)\nСКИБ: PDP", (232, 222, 248), (150, 115, 166)),
        (900, 145, 1250, 305, "CryptoContainerFacade\n+ seal(payload)\nFacade", (213, 232, 212), (130, 179, 102)),
        (70, 405, 420, 565, "CryptoContainerAdapter\n+ open(container)\nAdapter", (255, 242, 204), (214, 182, 86)),
        (485, 405, 835, 565, "Sanitizer\n+ clean(payload)\nСКИБ: очистка", (248, 206, 204), (184, 84, 80)),
        (900, 405, 1250, 565, "Kafka connect factory\n+ connect()\nFactory Method style", (225, 213, 231), (150, 115, 166)),
    ]
    for box in boxes:
        draw_box(draw, box[:4], box[4], box[5], box[6])
    draw_arrow(draw, (420, 225), (485, 225), "classifies")
    draw_arrow(draw, (835, 225), (900, 225), "authorizes")
    draw_arrow(draw, (420, 485), (485, 485), "opens then cleans")
    img.save(path)


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_report():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    architecture = ASSET_DIR / "architecture.png"
    sequence = ASSET_DIR / "sequence.png"
    classes = ASSET_DIR / "classes.png"
    architecture_png(architecture)
    sequence_png(sequence)
    classes_png(classes)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Проектная работа\nCrypto Broker Kafka")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph("Защищенный мини-мессенджер с отдельным потоком датчиков, криптоконтейнерами и политикой безопасности")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Цель", level=1)
    doc.add_paragraph(
        "Разработать прототип системы обмена сообщениями с криптографической защитой информации. "
        "Основной сценарий - защищенный мессенджер. Дополнительно выделен автоматический поток данных датчиков, "
        "который генерируется отдельным контейнером и подходит для визуализации в Grafana."
    )

    doc.add_heading("Функциональность", level=1)
    for item in [
        "веб-интерфейс на localhost:8088 для переписки в формате мессенджера;",
        "sensor-generator для автоматической отправки температуры и влажности;",
        "раздельные Kafka-топики для переписки и сенсоров;",
        "policy-engine для принятия решений allow/deny;",
        "crypto-gateway для упаковки payload в криптоконтейнер;",
        "filter для расшифрования, очистки и валидации;",
        "consumer и Grafana для просмотра результата.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Топики Kafka", level=1)
    add_table(
        doc,
        ["Поток", "Raw topic", "Crypto topic", "Filtered topic"],
        [
            ["Чат", "messages.raw", "messages.crypto", "messages.filtered"],
            ["Датчики", "sensors.raw", "sensors.crypto", "sensors.data.filtered"],
            ["Политика", "policy.requests", "policy.decisions", "-"],
        ],
    )

    doc.add_heading("Архитектура", level=1)
    doc.add_paragraph(
        "Система разделена на изолированные контейнеры. Web App, producer и sensor-generator не имеют ключей расшифрования. "
        "Policy-engine принимает решение безопасности, а crypto-gateway применяет это решение. "
        "Открытый payload не публикуется в crypto-топики."
    )
    doc.add_picture(str(architecture), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Диаграмма взаимодействия", level=1)
    doc.add_picture(str(sequence), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Диаграмма классов", level=1)
    doc.add_picture(str(classes), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Политика безопасности", level=1)
    doc.add_paragraph("Файл политики находится в lab-kafka-iot/policy/policy.json.")
    add_code(
        doc,
        """
        allowed_sources: web:chat, sensor-generator, producer:chat, producer:sensor, producer:/sample
        allowed_message_types: chat-message, sensor-data
        max_text_length: 240
        temperature: -20..50
        humidity: 0..100
        default_decision: deny
        """,
    )

    doc.add_heading("Криптоконтейнер", level=1)
    doc.add_paragraph(
        "После разрешения политики crypto-gateway создает контейнер с зашифрованным payload. "
        "В messages.crypto и sensors.crypto хранится ciphertext, а не открытый JSON."
    )
    add_code(
        doc,
        """
        {
          "container_version": "1.0",
          "message_type": "chat-message",
          "source_topic": "messages.raw",
          "algorithm": "Fernet(AES-128-CBC-HMAC-SHA256)",
          "ciphertext": "gAAAAAB..."
        }
        """,
    )

    doc.add_heading("Использованные шаблоны", level=1)
    add_table(
        doc,
        ["Шаблон", "Где используется", "Назначение"],
        [
            ["Strategy", "MessageTypeStrategy", "определение типа сообщения и маршрута"],
            ["Facade", "CryptoContainerFacade", "единый интерфейс создания криптоконтейнера"],
            ["Adapter", "CryptoContainerAdapter", "преобразование криптоконтейнера обратно в payload"],
            ["Factory Method style", "connect()", "создание Kafka-клиентов"],
            ["СКИБ: PDP/PEP", "policy-engine / crypto-gateway", "разделение принятия и применения решений"],
            ["СКИБ: очистка данных", "filter", "валидация и очистка перед итоговым топиком"],
        ],
    )

    doc.add_heading("Запуск", level=1)
    add_code(
        doc,
        """
        docker compose up --build -d kafka kafka-init sensor-generator policy-engine crypto-gateway filter consumer app grafana
        http://localhost:8088
        http://localhost:3001
        """,
    )

    doc.add_heading("Вывод", level=1)
    doc.add_paragraph(
        "Получен прототип кибериммунной системы обмена сообщениями: чат и автоматический поток датчиков разделены по топикам, "
        "сообщения проходят через политику безопасности, шифруются в криптоконтейнеры и публикуются в итоговые "
        "топики только после расшифрования, очистки и валидации."
    )

    doc.add_heading("Источники", level=1)
    doc.add_paragraph("https://securitybydesign.ru/templates/")
    doc.add_paragraph("https://refactoringu.ru/ru/design-patterns/catalog.html")

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
